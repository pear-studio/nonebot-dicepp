import json
import os
import re
import time
import uuid
import zlib
from typing import Any, Dict, List, Optional, Tuple

try:
    import requests  # type: ignore
except Exception:  # pragma: no cover
    requests = None

from core.bot import Bot
from core.data import DataManagerError, DataChunkBase, custom_data_chunk
from core.config import CFG_MASTER
from core.command.const import *
from core.command import BotCommandBase, BotSendFileCommand, BotSendMsgCommand
from core.command import UserCommandBase, custom_user_command
from core.communication import GroupMessagePort, MessageMetaData
from utils.time import get_current_date_str, str_to_datetime
from utils.logger import dice_log

# 日志数据库后端（将记录存入 SQLite，导出从 DB 读取）
try:
    from .log_db import (
        get_connection,
        upsert_log,
        insert_record,
        fetch_records,
        delete_log,
        delete_records_by_message_id,
        set_recording,
        update_log_upload,
    )
except Exception:
    # 兼容导入失败场景，保持旧逻辑可运行（但不会用到 DB）
    get_connection = None  # type: ignore
    upsert_log = None  # type: ignore
    insert_record = None  # type: ignore
    fetch_records = None  # type: ignore
    delete_log = None  # type: ignore
    delete_records_by_message_id = None  # type: ignore
    set_recording = None  # type: ignore
    update_log_upload = None  # type: ignore

# 旧版本使用的常量，保留以兼容外部引用或进行数据迁移
DC_LOG_SESSION = "log_session"
DCK_ACTIVE = "active"
DCK_START_TIME = "start_time"
DCK_RECORDS = "records"
DCK_COLOR_MAP = "color_map"
DCK_MSG_COUNT = "msg_count"
DCK_LAST_HOUR_WARN = "last_hour_warn"
DCK_FILTER_OUTSIDE = "filter_outside"
DCK_FILTER_COMMAND = "filter_command"
DCK_FILTER_BOT = "filter_bot"
DCK_FILTER_MEDIA = "filter_media"
DCK_FILTER_FORUM_CODE = "filter_forum_code"

LEGACY_KEYS = {
    DCK_ACTIVE,
    DCK_START_TIME,
    DCK_RECORDS,
    DCK_COLOR_MAP,
    DCK_MSG_COUNT,
    DCK_LAST_HOUR_WARN,
    DCK_FILTER_OUTSIDE,
    DCK_FILTER_COMMAND,
    DCK_FILTER_BOT,
    DCK_FILTER_MEDIA,
    DCK_FILTER_FORUM_CODE,
}

# 新版日志存储结构常量
LOG_GROUP_CURRENT = "current"
LOG_GROUP_LOGS = "logs"
LOG_GROUP_FILTERS = "filters"
LOG_GROUP_NAME_INDEX = "name_index"

FILTER_OUTSIDE = "outside"
FILTER_COMMAND = "command"
FILTER_BOT = "bot"
FILTER_MEDIA = "media"
FILTER_FORUM_CODE = "forum_code"

LOG_KEY_NAME = "name"
LOG_KEY_CREATED_AT = "created_at"
LOG_KEY_UPDATED_AT = "updated_at"
LOG_KEY_RECORDING = "recording"
LOG_KEY_RECORDS = "records"
LOG_KEY_COLOR_MAP = "color_map"
LOG_KEY_STATS = "stats"
LOG_KEY_SESSION_COUNT = "session_count"
LOG_KEY_RECORD_BEGIN_AT = "record_begin_at"
LOG_KEY_LAST_WARN = "last_warn"
LOG_KEY_UPLOAD = "upload"
LOG_KEY_UPLOAD_TIME = "time"
LOG_KEY_UPLOAD_FILE = "file"
LOG_KEY_UPLOAD_NOTE = "note"
LOG_KEY_SOURCE = "source"

DEFAULT_FILTERS = {
    FILTER_OUTSIDE: False,
    FILTER_COMMAND: False,
    FILTER_BOT: False,
    FILTER_MEDIA: False,
    FILTER_FORUM_CODE: False,
}

ROLL_CONTEXT_KEYWORDS = ["掷骰", "检定", "攻击", "豁免", "暗骰", "命中", "伤害"]
ROLL_SUCCESS_KEYWORDS = ["成功", "命中"]
ROLL_FAILURE_KEYWORDS = ["失败", "未命中", "落败", "扑空"]

RE_ATTR_ARROW = re.compile(
    r"(力量|敏捷|体质|智力|感知|魅力|体型|外貌|意志|教育|幸运|SAN值?|SAN|理智|HP|生命|体力)\s*[:：]?\s*(\d+)\s*(?:->|→|=>|到|至)\s*(\d+)",
    re.IGNORECASE,
)
RE_ATTR_DELTA = re.compile(
    r"(力量|敏捷|体质|智力|感知|魅力|体型|外貌|意志|教育|幸运|SAN值?|SAN|理智|HP|生命|体力)\s*[:：]?\s*([+-]\d+)",
    re.IGNORECASE,
)
RE_DICE_RESULT = re.compile(r"(\d+)d(\d+)\s*=\s*(\d+)")
RE_ROLLER_PREFIX = re.compile(r"^([^\s：:\n]+)")

ATTR_NAME_ALIASES = {
    "SAN": "SAN",
    "SAN值": "SAN",
    "SANVALUE": "SAN",
    "理智": "SAN",
    "HP": "HP",
    "HP值": "HP",
    "HPVALUE": "HP",
    "体力": "HP",
    "生命": "HP",
}

COLOR_POOL = [
    "FF0000", "1E90FF", "228B22", "FF8C00", "9400D3", "DC143C", "20B2AA", "8B4513",
    "FF1493", "2E8B57", "4169E1", "DAA520", "C71585", "008B8B", "B03060", "556B2F",
]

UPLOAD_ENDPOINT_DEFAULT = "https://dice.weizaima.com/dice/api/log"
UPLOAD_VERSION = 105
CFG_LOG_UPLOAD_ENABLE = "log_upload_enable"
CFG_LOG_UPLOAD_ENDPOINT = "log_upload_endpoint"
CFG_LOG_UPLOAD_TOKEN = "log_upload_token"
# 限制单个日志在内存中保留的最大记录条数（可通过配置覆盖）。
CFG_LOG_MAX_RECORDS = "log_max_records"
LOG_MAX_RECORDS_DEFAULT = 5000
# 在启用数据库存储后，内存侧仅保留少量最新记录作为保险，避免深拷贝过大数据。
LOG_IN_MEMORY_SAFE_LIMIT = 50
# 内存中 stats.participants 字典最大条目数，超出将保留最活跃的用户
LOG_PARTICIPANTS_LIMIT = 500
# 内存中 color_map 字典最大条目数
LOG_COLOR_MAP_LIMIT = 500
# 每个骰面类型在 dice_faces.users 中保留的最大用户数
LOG_DICE_USERS_LIMIT = 100


def _pick_color(color_map: Dict[str, str], user_id: str) -> str:
    if user_id not in color_map:
        color_map[user_id] = COLOR_POOL[len(color_map) % len(COLOR_POOL)]
    return color_map[user_id]


@custom_data_chunk(identifier=DC_LOG_SESSION)
class _(DataChunkBase):  # noqa: E742
    def __init__(self):
        super().__init__()


def _now_str() -> str:
    return get_current_date_str()


def _sanitize_filename(name: str) -> str:
    cleaned = re.sub(r"[\\/:*?\"<>|]", "_", name).strip()
    return cleaned or "log"


def _normalize_attr_name(raw: str) -> str:
    trimmed = (raw or "").strip()
    key = trimmed.upper()
    return ATTR_NAME_ALIASES.get(key, trimmed)


def _empty_stats() -> Dict[str, Any]:
    return {
        "messages": 0,
        "participants": {},  # user_id -> {"count": int, "nickname": Optional[str]}
        "rolls": {
            "success": 0,
            "failure": 0,
            "critical_success": 0,
            "critical_failure": 0,
        },
        "attributes": {},
        "dice_faces": {},  # face -> {sum: float, count: int, users: {id: {sum: float, count: int, nickname: str}}}
    }


def _detect_roll_result(content: str) -> Optional[str]:
    if not content:
        return None
    flattened = content.replace("\n", " ")
    if not any(keyword in flattened for keyword in ROLL_CONTEXT_KEYWORDS):
        return None
    if "大成功" in flattened:
        return "critical_success"
    if "大失败" in flattened:
        return "critical_failure"
    success = any(keyword in flattened for keyword in ROLL_SUCCESS_KEYWORDS)
    failure = any(keyword in flattened for keyword in ROLL_FAILURE_KEYWORDS)
    if success and not failure:
        return "success"
    if failure and not success:
        return "failure"
    return None


def _detect_attr_changes(content: str) -> Dict[str, int]:
    changes: Dict[str, int] = {}
    if not content:
        return changes
    for match in RE_ATTR_ARROW.finditer(content):
        attr = _normalize_attr_name(match.group(1))
        delta = int(match.group(3)) - int(match.group(2))
        if delta:
            changes[attr] = changes.get(attr, 0) + delta
    for match in RE_ATTR_DELTA.finditer(content):
        attr = _normalize_attr_name(match.group(1))
        delta = int(match.group(2))
        if delta:
            changes[attr] = changes.get(attr, 0) + delta
    return changes


def _accumulate_roll_detail(stats: Dict[str, Any], record: Dict[str, Any]) -> None:
    content = record.get("content", "")
    if not content:
        return
    matches = list(RE_DICE_RESULT.finditer(content))
    if not matches:
        return

    participants = stats.setdefault("participants", {})
    dice_faces_stats = stats.setdefault("dice_faces", {})

    # 估算掷骰者昵称
    candidate_name = record.get("nickname") or ""
    stripped_content = content.strip()
    prefix_match = RE_ROLLER_PREFIX.match(stripped_content)
    if prefix_match:
        candidate_name = prefix_match.group(1)

    user_key = None
    display_name = candidate_name
    if candidate_name:
        for uid, info in participants.items():
            nickname = info.get("nickname") or uid
            if candidate_name == nickname or candidate_name == uid:
                user_key = uid
                display_name = nickname
                break
    if not user_key:
        user_key = f"name:{candidate_name}" if candidate_name else "unknown"
        if user_key not in participants and candidate_name:
            participants.setdefault(user_key, {"count": 0, "nickname": candidate_name})

    for match in matches:
        count = max(int(match.group(1)), 1)
        faces = int(match.group(2))
        result = int(match.group(3))
        if faces <= 1:
            continue
        max_val = count * faces
        norm = min(max(result, 0), max_val) / max_val

        face_entry = dice_faces_stats.setdefault(faces, {"sum": 0.0, "count": 0, "users": {}})
        face_entry["sum"] += norm
        face_entry["count"] += 1

        user_entry = face_entry["users"].setdefault(user_key, {"sum": 0.0, "count": 0, "nickname": display_name or user_key})
        user_entry["sum"] += norm
        user_entry["count"] += 1


def _update_stats_with_record(stats: Dict[str, Any], record: Dict[str, Any], *, source_is_bot: bool) -> None:
    user_id = record.get("user_id", "?")
    nickname = record.get("nickname")
    stats["messages"] = stats.get("messages", 0) + 1

    participants = stats.setdefault("participants", {})
    info = participants.setdefault(user_id, {"count": 0, "nickname": None})
    info["count"] = info.get("count", 0) + 1
    if nickname and nickname not in ("UNDEF_NAME", "----"):
        info["nickname"] = nickname

    if source_is_bot:
        result = _detect_roll_result(record.get("content", ""))
        if result:
            stats.setdefault("rolls", {}).setdefault(result, 0)
            stats["rolls"][result] += 1
        attr_changes = _detect_attr_changes(record.get("content", ""))
        if attr_changes:
            attr_stats = stats.setdefault("attributes", {})
            for attr, delta in attr_changes.items():
                attr_stats[attr] = attr_stats.get(attr, 0) + delta
        _accumulate_roll_detail(stats, record)


def _compute_log_stats(bot: Bot, records: List[Dict[str, Any]]) -> Dict[str, Any]:
    stats = _empty_stats()
    for rec in records:
        source = rec.get(LOG_KEY_SOURCE)
        if source is None:
            source_is_bot = rec.get("user_id") == bot.account
        else:
            source_is_bot = source == "bot"
        _update_stats_with_record(stats, rec, source_is_bot=source_is_bot)
    return stats


def _ensure_filters(payload: Dict[str, Any]) -> Dict[str, bool]:
    filters = payload.setdefault(LOG_GROUP_FILTERS, {})
    for key, default in DEFAULT_FILTERS.items():
        filters.setdefault(key, default)
    return filters


def _generate_log_id() -> str:
    return uuid.uuid4().hex


def _init_group_payload(bot: Bot, group_id: str, legacy: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    payload: Dict[str, Any] = {
        LOG_GROUP_CURRENT: "",
        LOG_GROUP_LOGS: {},
        LOG_GROUP_FILTERS: DEFAULT_FILTERS.copy(),
        LOG_GROUP_NAME_INDEX: {},
    }

    if isinstance(legacy, dict) and any(key in legacy for key in LEGACY_KEYS):
        filters = payload[LOG_GROUP_FILTERS]
        filters[FILTER_OUTSIDE] = legacy.get(DCK_FILTER_OUTSIDE, False)
        filters[FILTER_COMMAND] = legacy.get(DCK_FILTER_COMMAND, False)
        filters[FILTER_BOT] = legacy.get(DCK_FILTER_BOT, False)
        filters[FILTER_MEDIA] = legacy.get(DCK_FILTER_MEDIA, False)
        filters[FILTER_FORUM_CODE] = legacy.get(DCK_FILTER_FORUM_CODE, False)

        legacy_records = legacy.get(DCK_RECORDS, []) or []
        for rec in legacy_records:
            rec.setdefault(LOG_KEY_SOURCE, "bot" if rec.get("user_id") == bot.account else "user")
        if legacy_records:
            log_name = legacy.get("log_name") or legacy.get(LOG_KEY_NAME) or "默认日志"
            log_id = _generate_log_id()
            start_time = legacy.get(DCK_START_TIME, _now_str())
            log_entry = {
                LOG_KEY_NAME: log_name,
                LOG_KEY_CREATED_AT: start_time,
                LOG_KEY_UPDATED_AT: legacy.get(DCK_LAST_HOUR_WARN, start_time),
                LOG_KEY_RECORDING: legacy.get(DCK_ACTIVE, False),
                LOG_KEY_RECORDS: legacy_records,
                LOG_KEY_COLOR_MAP: legacy.get(DCK_COLOR_MAP, {}),
                LOG_KEY_STATS: _compute_log_stats(bot, legacy_records),
                LOG_KEY_SESSION_COUNT: legacy.get(DCK_MSG_COUNT, 0),
                LOG_KEY_RECORD_BEGIN_AT: legacy.get(DCK_START_TIME, start_time),
                LOG_KEY_LAST_WARN: legacy.get(DCK_LAST_HOUR_WARN, start_time),
                LOG_KEY_UPLOAD: {},
            }
            payload[LOG_GROUP_LOGS][log_id] = log_entry
            payload[LOG_GROUP_NAME_INDEX][log_name.lower()] = log_id
            if legacy.get(DCK_ACTIVE, False):
                payload[LOG_GROUP_CURRENT] = log_id

    bot.data_manager.set_data(DC_LOG_SESSION, [group_id], payload)
    return payload


def _rebuild_name_index(payload: Dict[str, Any]) -> None:
    name_index = payload.setdefault(LOG_GROUP_NAME_INDEX, {})
    name_index.clear()
    for log_id, entry in payload.get(LOG_GROUP_LOGS, {}).items():
        log_name = entry.get(LOG_KEY_NAME, log_id)
        key = log_name.lower()
        if key in name_index and name_index[key] != log_id:
            key = f"{key}_{log_id[:6]}"
        name_index[key] = log_id


def _load_group_payload(bot: Bot, group_id: str) -> Dict[str, Any]:
    try:
        payload = bot.data_manager.get_data(DC_LOG_SESSION, [group_id])
    except DataManagerError:
        payload = None

    if not isinstance(payload, dict) or LOG_GROUP_LOGS not in payload:
        payload = _init_group_payload(bot, group_id, payload if isinstance(payload, dict) else None)
    else:
        mutated = False
        filters = _ensure_filters(payload)
        if filters.keys() != DEFAULT_FILTERS.keys():
            mutated = True
        logs = payload.setdefault(LOG_GROUP_LOGS, {})
        for log_id, entry in list(logs.items()):
            records = entry.setdefault(LOG_KEY_RECORDS, [])
            for rec in records:
                rec.setdefault(LOG_KEY_SOURCE, "bot" if rec.get("user_id") == bot.account else "user")
            if LOG_KEY_STATS not in entry:
                entry[LOG_KEY_STATS] = _compute_log_stats(bot, records)
                mutated = True
            else:
                stats = entry[LOG_KEY_STATS]
                if "dice_faces" not in stats:
                    stats["dice_faces"] = {}
                    entry[LOG_KEY_STATS] = stats
                    mutated = True
            if LOG_KEY_COLOR_MAP not in entry:
                entry[LOG_KEY_COLOR_MAP] = {}
                mutated = True
            entry.setdefault(LOG_KEY_CREATED_AT, _now_str())
            entry.setdefault(LOG_KEY_UPDATED_AT, entry.get(LOG_KEY_CREATED_AT))
            entry.setdefault(LOG_KEY_RECORDING, False)
            entry.setdefault(LOG_KEY_SESSION_COUNT, 0)
            entry.setdefault(LOG_KEY_RECORD_BEGIN_AT, entry.get(LOG_KEY_UPDATED_AT))
            entry.setdefault(LOG_KEY_LAST_WARN, entry.get(LOG_KEY_RECORD_BEGIN_AT))
            entry.setdefault(LOG_KEY_UPLOAD, {})
            logs[log_id] = entry
        if LOG_GROUP_CURRENT not in payload:
            payload[LOG_GROUP_CURRENT] = ""
            mutated = True
        _rebuild_name_index(payload)
        if mutated:
            bot.data_manager.set_data(DC_LOG_SESSION, [group_id], payload)
    return payload


def _save_group_payload(bot: Bot, group_id: str, payload: Dict[str, Any]) -> None:
    bot.data_manager.set_data(DC_LOG_SESSION, [group_id], payload)


def _find_log_id_by_name(payload: Dict[str, Any], name: str) -> Optional[str]:
    if not name:
        return None
    idx = payload.get(LOG_GROUP_NAME_INDEX, {})
    key = name.lower()
    if key in idx:
        return idx[key]
    for log_id, entry in payload.get(LOG_GROUP_LOGS, {}).items():
        if entry.get(LOG_KEY_NAME, "").lower() == key:
            return log_id
    return None


def _append_record_to_entry(log_entry: Dict[str, Any], record: Dict[str, Any], *, source_is_bot: bool) -> None:
    records = log_entry.setdefault(LOG_KEY_RECORDS, [])
    records.append(record)
    color_map = log_entry.setdefault(LOG_KEY_COLOR_MAP, {})
    _pick_color(color_map, record.get("user_id", ""))
    stats = log_entry.setdefault(LOG_KEY_STATS, _empty_stats())
    stats.setdefault("dice_faces", {})
    _update_stats_with_record(stats, record, source_is_bot=source_is_bot)
    log_entry[LOG_KEY_UPDATED_AT] = record.get("time", _now_str())


def _append_record_to_db(group_id: str, log_id: str, log_entry: Dict[str, Any], record: Dict[str, Any], *, source_is_bot: bool) -> None:
    """将记录写入数据库，同时在内存里仅维护必要的统计与配色，避免内存暴涨。"""
    # 1) 确保日志元数据存在（旧日志可能在 DB 中尚未建档）
    if get_connection and upsert_log:
        try:
            conn = get_connection()
            try:
                # 取 filters 快照，便于后续导出/查看
                filters = log_entry.get(LOG_KEY_UPLOAD)  # dummy to satisfy type checker
            except Exception:
                filters = None
            finally:
                # 重置，真实 filters 从调用方 payload 中难以直接传入，这里尽量填充已有字段
                pass
            try:
                upsert_log(conn, {
                    "id": log_id,
                    "group_id": group_id,
                    "name": log_entry.get(LOG_KEY_NAME, log_id),
                    "created_at": log_entry.get(LOG_KEY_CREATED_AT, record.get("time", _now_str())),
                    "updated_at": record.get("time", _now_str()),
                    "recording": True if log_entry.get(LOG_KEY_RECORDING) else False,
                    "record_begin_at": log_entry.get(LOG_KEY_RECORD_BEGIN_AT, record.get("time", _now_str())),
                    "last_warn": log_entry.get(LOG_KEY_LAST_WARN, log_entry.get(LOG_KEY_RECORD_BEGIN_AT, record.get("time", _now_str()))),
                    "filter_outside": 0,
                    "filter_command": 0,
                    "filter_bot": 0,
                    "filter_media": 0,
                    "filter_forum_code": 0,
                    "upload_time": log_entry.get(LOG_KEY_UPLOAD, {}).get(LOG_KEY_UPLOAD_TIME),
                    "upload_file": log_entry.get(LOG_KEY_UPLOAD, {}).get(LOG_KEY_UPLOAD_FILE),
                    "upload_note": log_entry.get(LOG_KEY_UPLOAD, {}).get(LOG_KEY_UPLOAD_NOTE),
                    "url": log_entry.get(LOG_KEY_UPLOAD, {}).get("url"),
                })
                conn.commit()
            finally:
                conn.close()
        except Exception as e:
            dice_log(f"[LogDB] upsert before insert error: {e}")

    # 2) 写入记录
    if get_connection and insert_record:
        try:
            conn = get_connection()
            try:
                insert_record(
                    conn,
                    log_id,
                    time=record.get("time", _now_str()),
                    user_id=str(record.get("user_id") or ""),
                    nickname=record.get("nickname") or str(record.get("user_id") or ""),
                    content=record.get("content", ""),
                    source=record.get(LOG_KEY_SOURCE, "user"),
                    message_id=record.get("message_id"),
                )
                conn.commit()
            finally:
                conn.close()
        except Exception as e:
            dice_log(f"[LogDB] insert_record error: {e}")

    # 3) 内存：只维护统计与颜色映射
    color_map = log_entry.setdefault(LOG_KEY_COLOR_MAP, {})
    _pick_color(color_map, record.get("user_id", ""))
    stats = log_entry.setdefault(LOG_KEY_STATS, _empty_stats())
    stats.setdefault("dice_faces", {})
    _update_stats_with_record(stats, record, source_is_bot=source_is_bot)
    log_entry[LOG_KEY_UPDATED_AT] = record.get("time", _now_str())


def _trim_records_if_needed(bot: Bot, entry: Dict[str, Any]) -> None:
    """根据配置裁剪过多的历史记录，避免内存无限增长。
    - 读取配置 CFG_LOG_MAX_RECORDS（默认 5000）。
    - 当超限时，删除最早的一批（回落到阈值的 80%）以减少频繁切片的成本。
    """
    try:
        cfg_val = bot.cfg_helper.get_config(CFG_LOG_MAX_RECORDS)[0]
        limit = int(str(cfg_val).strip())
    except Exception:
        limit = LOG_MAX_RECORDS_DEFAULT
    # 强制收敛至安全上限，确保 DataManager 深拷贝时不会携带过多记录。
    if LOG_IN_MEMORY_SAFE_LIMIT > 0:
        if limit <= 0 or limit > LOG_IN_MEMORY_SAFE_LIMIT:
            limit = LOG_IN_MEMORY_SAFE_LIMIT
    if limit is None or limit <= 0:
        return
    records = entry.get(LOG_KEY_RECORDS, [])
    cur_len = len(records)
    if cur_len > limit:
        # 回落到 80% 的水位，减少切片频率
        target = max(int(limit * 0.8), 1)
        drop_n = cur_len - target
        if drop_n > 0:
            del records[:drop_n]
            entry[LOG_KEY_RECORDS] = records


def _trim_stats_if_needed(entry: Dict[str, Any]) -> None:
    """裁剪 stats 中的 participants 和 dice_faces.users，避免无限增长。"""
    stats = entry.get(LOG_KEY_STATS)
    if not isinstance(stats, dict):
        return

    # 1) 裁剪 participants：保留消息数最多的用户
    participants = stats.get("participants")
    if isinstance(participants, dict) and len(participants) > LOG_PARTICIPANTS_LIMIT:
        sorted_items = sorted(
            participants.items(),
            key=lambda x: x[1].get("count", 0) if isinstance(x[1], dict) else 0,
            reverse=True
        )[:LOG_PARTICIPANTS_LIMIT]
        stats["participants"] = dict(sorted_items)

    # 2) 裁剪 dice_faces.users：每个骰面保留掷骰次数最多的用户
    dice_faces = stats.get("dice_faces")
    if isinstance(dice_faces, dict):
        for face, face_info in dice_faces.items():
            if not isinstance(face_info, dict):
                continue
            users = face_info.get("users")
            if isinstance(users, dict) and len(users) > LOG_DICE_USERS_LIMIT:
                sorted_users = sorted(
                    users.items(),
                    key=lambda x: x[1].get("count", 0) if isinstance(x[1], dict) else 0,
                    reverse=True
                )[:LOG_DICE_USERS_LIMIT]
                face_info["users"] = dict(sorted_users)


def _trim_color_map_if_needed(entry: Dict[str, Any]) -> None:
    """裁剪 color_map，避免无限增长。保留活跃用户的颜色映射。"""
    color_map = entry.get(LOG_KEY_COLOR_MAP)
    if not isinstance(color_map, dict) or len(color_map) <= LOG_COLOR_MAP_LIMIT:
        return

    # 优先保留 participants 中存在的用户
    stats = entry.get(LOG_KEY_STATS, {})
    participants = stats.get("participants", {}) if isinstance(stats, dict) else {}
    active_users = set(participants.keys())

    new_map: Dict[str, str] = {}
    # 先添加活跃用户
    for uid in active_users:
        if uid in color_map and len(new_map) < LOG_COLOR_MAP_LIMIT:
            new_map[uid] = color_map[uid]
    # 再填充其他（按原顺序）
    for uid, color in color_map.items():
        if uid not in new_map and len(new_map) < LOG_COLOR_MAP_LIMIT:
            new_map[uid] = color

    entry[LOG_KEY_COLOR_MAP] = new_map


def _should_filter(filters: Dict[str, bool], content: str, *, is_bot: bool) -> bool:
    text = (content or "").strip()
    if filters.get(FILTER_OUTSIDE) and (
        (text.startswith('(') and text.endswith(')')) or
        (text.startswith('（') and text.endswith('）'))
    ):
        return True
    if filters.get(FILTER_COMMAND) and (text.startswith('.') or text.startswith('。')):
        return True
    if filters.get(FILTER_BOT) and is_bot:
        return True
    if filters.get(FILTER_MEDIA):
        lowered = text.lower()
        if any(token in lowered for token in ("[cq:image", "[cq:face", "[cq:emoji", "[cq:video")):
            return True
    if "[CQ:file," in content:
        return True
    return False


def should_filter_record(bot: Bot, group_id: str, user_id: str, content: str, is_bot: bool = False) -> bool:
    payload = _load_group_payload(bot, group_id)
    filters = _ensure_filters(payload)
    return _should_filter(filters, content, is_bot=is_bot)


def _generate_forum_code_log(records: List[Dict[str, Any]]) -> str:
    forum_code = []
    for record in records:
        time = record.get("time", "未知时间")
        nickname = record.get("nickname", "未知用户")
        content = record.get("content", "")
        forum_code.append(f"[color=#9ca3af]{time}[/color][color=#f99252] <{nickname}>{content} [/color]")
    return "\n".join(forum_code)


def append_log_record(bot: Bot, group_id: str, user_id: str, nickname: str, content: str,
                      message_id: Optional[str] = None):
    try:
        record_incoming_message(
            bot,
            group_id,
            str(user_id or ""),
            nickname or ("骰娘" if user_id == bot.account else user_id),
            content,
            message_id,
            is_bot=True,
        )
    except Exception:
        pass


class _StatsFormatter:
    @staticmethod
    def format(log_entry: Dict[str, Any]) -> str:
        name = log_entry.get(LOG_KEY_NAME, "?")
        stats = log_entry.get(LOG_KEY_STATS, _empty_stats())
        messages = stats.get("messages", 0)
        participants = stats.get("participants", {})
        rolls = stats.get("rolls", {})
        attributes = stats.get("attributes", {})
        dice_faces = stats.get("dice_faces", {})

        lines = [f"日志《{name}》统计", f"- 记录消息：{messages} 条"]
        lines.append(f"- 参与者：{len(participants)} 人")
        if participants:
            top = sorted(
                participants.items(),
                key=lambda item: item[1].get("count", 0),
                reverse=True,
            )[:5]
            top_str = "、".join(
                f"{info.get('nickname') or uid}({info.get('count', 0)})"
                for uid, info in top
            )
            lines.append(f"- TOP5 活跃：{top_str}")
        else:
            lines.append("- TOP5 活跃：暂无数据")

        success = rolls.get("success", 0)
        failure = rolls.get("failure", 0)
        c_success = rolls.get("critical_success", 0)
        c_failure = rolls.get("critical_failure", 0)
        lines.append(
            f"- 检定统计：成功 {success}，失败 {failure}，大成功 {c_success}，大失败 {c_failure}"
        )
        if attributes:
            attr_line = "，".join(
                f"{attr}{delta:+d}" for attr, delta in attributes.items() if delta
            )
            lines.append(f"- 属性变化：{attr_line}")
        else:
            lines.append("- 属性变化：暂无记录")

        if dice_faces:
            lines.append("- 骰运统计：")
            for face in sorted(dice_faces.keys()):
                face_info = dice_faces[face]
                total_count = face_info.get("count", 0)
                total_sum = face_info.get("sum", 0.0)
                if not total_count:
                    continue
                avg_percent = total_sum / total_count * 100
                user_entries = face_info.get("users", {})
                top_users = []
                for uid, info in user_entries.items():
                    cnt = info.get("count", 0)
                    if not cnt:
                        continue
                    ratio = info.get("sum", 0.0) / cnt
                    nickname = info.get("nickname") or uid
                    top_users.append((ratio, nickname))
                top_users.sort(key=lambda x: x[0], reverse=True)
                top_str = "、".join(
                    f"{name}({ratio*100:.1f}%)" for ratio, name in top_users[:3]
                ) if top_users else "暂无数据"
                lines.append(f"  · d{face} 平均 {avg_percent:.1f}% ，运气TOP：{top_str}")

        lines.append(f"- 最近更新时间：{log_entry.get(LOG_KEY_UPDATED_AT, '-')}")
        return "\n".join(lines)


# 旧的开关提示文本
LOC_LOG_ON_START = "log_on_start"
LOC_LOG_ON_ALREADY = "log_on_already"
LOC_LOG_OFF_NOT_ACTIVE = "log_off_not_active"
LOC_LOG_OFF_RESULT = "log_off_result"
LOC_LOG_USAGE = "log_usage"
LOC_LOG_SET_MENU = "log_set_menu"
LOC_LOG_SET_TOGGLED = "log_set_toggled"
LOC_LOG_HELP = "log_help"
LOC_LOG_FOLDER_FAIL = "log_folder_fail"
LOC_LOG_FOLDER_HINT = "log_folder_hint"
LOC_LOG_FOLDER_POLICY = "log_folder_policy"


class _LogMessages:
    def __init__(self):
        self.new_started = "日志《{name}》已创建并开始记录。"
        self.already_running = "日志《{name}》已经在记录中。"
        self.resume = "日志《{name}》继续记录。"
        self.paused = "日志《{name}》已暂停，可使用 .log on 继续。"
        self.no_current = "当前没有正在使用的日志。"
        self.no_target = "未找到名为《{name}》的日志。"
        self.need_empty_current = "请先使用 .log halt 或 .log end 结束当前日志，再执行该操作。"
        self.list_header = "日志列表："
        self.deleted = "日志《{name}》已删除。"
        self.halted = "日志《{name}》已停止记录。"
        self.end_summary = (
            "日志《{name}》已导出，共 {count} 条消息。"
            "已向群文件上传生成的日志（如存在“跑团log”文件夹将自动归档）。"
        )
        self.pause_before_delete = "日志《{name}》正在记录，请先 .log off 或 .log halt。"
        self.not_recording = "日志《{name}》当前处于暂停状态。"
        self.switch_success = "已切换至日志《{name}》。"


class _LogHelper:
    def __init__(self, bot: Bot):
        self.bot = bot

    def load_payload(self, group_id: str) -> Dict[str, Any]:
        return _load_group_payload(self.bot, group_id)

    def save_payload(self, group_id: str, payload: Dict[str, Any]) -> None:
        _save_group_payload(self.bot, group_id, payload)

    def ensure_log_available(self, payload: Dict[str, Any], log_id: str) -> Optional[Dict[str, Any]]:
        return payload.get(LOG_GROUP_LOGS, {}).get(log_id)

    def make_record(self, user_id: str, nickname: str, content: str, message_id: Optional[str], source: str) -> Dict[str, Any]:
        record = {
            "time": _now_str(),
            "user_id": user_id,
            "nickname": nickname or user_id,
            "content": content,
            LOG_KEY_SOURCE: source,
        }
        if message_id:
            record["message_id"] = message_id
        return record


class _LogFormatter:
    @staticmethod
    def status_line(log_id: str, entry: Dict[str, Any], current_id: str) -> str:
        name = entry.get(LOG_KEY_NAME, log_id)
        stats = entry.get(LOG_KEY_STATS, {})
        count = stats.get("messages", len(entry.get(LOG_KEY_RECORDS, [])))
        created = entry.get(LOG_KEY_CREATED_AT, "-")
        upload = entry.get(LOG_KEY_UPLOAD, {}).get(LOG_KEY_UPLOAD_TIME)
        if log_id == current_id:
            if entry.get(LOG_KEY_RECORDING):
                status_flag = "▶"
                status_text = "进行中"
            else:
                status_flag = "▍"
                status_text = "已暂停"
        else:
            status_flag = "•"
            status_text = "未启用"
        parts = [f"{status_flag}《{name}》", f"状态:{status_text}", f"记录:{count} 条", f"创建:{created}"]
        if upload:
            parts.append(f"上次导出:{upload}")
        return " | ".join(parts)


class _Reminder:
    @staticmethod
    def should_notify_session(session_count: int) -> bool:
        return session_count > 0 and session_count % 1000 == 0

    @staticmethod
    def should_notify_hour(begin_at: str, last_warn_at: str, now: str) -> bool:
        try:
            begin_dt = str_to_datetime(begin_at)
            last_dt = str_to_datetime(last_warn_at)
            now_dt = str_to_datetime(now)
        except Exception:
            return False
        if begin_dt is None or last_dt is None or now_dt is None:
            return False
        elapsed_hours = int((now_dt - begin_dt).total_seconds() // 3600)
        warned_hours = int((last_dt - begin_dt).total_seconds() // 3600)
        return elapsed_hours >= 2 and (elapsed_hours - warned_hours) >= 2


def record_incoming_message(bot: Bot,
                            group_id: str,
                            user_id: str,
                            nickname: str,
                            content: str,
                            message_id: Optional[str],
                            *,
                            is_bot: bool) -> List[BotCommandBase]:
    if not group_id:
        return []

    payload = _load_group_payload(bot, group_id)
    current_id = payload.get(LOG_GROUP_CURRENT, "")
    if not current_id:
        return []

    logs = payload.get(LOG_GROUP_LOGS, {})
    entry = logs.get(current_id)
    if not entry or not entry.get(LOG_KEY_RECORDING):
        return []

    filters = _ensure_filters(payload)
    if _should_filter(filters, content, is_bot=is_bot):
        return []

    record = {
        "time": _now_str(),
        "user_id": str(user_id or ""),
        "nickname": nickname or ("骰娘" if str(user_id) == bot.account else user_id),
        "content": content,
        LOG_KEY_SOURCE: "bot" if is_bot else "user",
    }
    if message_id:
        record["message_id"] = message_id

    commands: List[BotCommandBase] = []
    if not is_bot:
        session_count = entry.get(LOG_KEY_SESSION_COUNT, 0) + 1
        entry[LOG_KEY_SESSION_COUNT] = session_count
        now_time = record["time"]
        if _Reminder.should_notify_session(session_count):
            commands.append(BotSendMsgCommand(
                bot.account,
                f"（日志《{entry.get(LOG_KEY_NAME, current_id)}》已记录 {session_count} 条消息。）",
                [GroupMessagePort(group_id)],
            ))
        begin_at = entry.get(LOG_KEY_RECORD_BEGIN_AT, now_time)
        last_warn = entry.get(LOG_KEY_LAST_WARN, begin_at)
        if entry.get(LOG_KEY_RECORDING) and _Reminder.should_notify_hour(begin_at, last_warn, now_time):
            try:
                hours = int((str_to_datetime(now_time) - str_to_datetime(begin_at)).total_seconds() // 3600)
            except Exception:
                hours = 1
            commands.append(BotSendMsgCommand(
                bot.account,
                f"（日志《{entry.get(LOG_KEY_NAME, current_id)}》已连续记录超过 {hours} 小时，如果已经整理完毕请使用 .log end 导出。）",
                [GroupMessagePort(group_id)],
            ))
            entry[LOG_KEY_LAST_WARN] = now_time

    _append_record_to_db(group_id, current_id, entry, record, source_is_bot=is_bot)
    # 不再堆积内存 records，仅保留统计；裁剪留作安全网（不会影响）
    _trim_records_if_needed(bot, entry)
    # 裁剪 stats 和 color_map，防止无限增长
    _trim_stats_if_needed(entry)
    _trim_color_map_if_needed(entry)
    payload[LOG_GROUP_LOGS][current_id] = entry
    _save_group_payload(bot, group_id, payload)
    return commands


@custom_user_command(readable_name="跑团日志指令",
                     priority=DPP_COMMAND_PRIORITY_DEFAULT,
                     flag=DPP_COMMAND_FLAG_DEFAULT,
                     cluster=DPP_COMMAND_CLUSTER_DEFAULT,
                     group_only=True)
class LogCommand(UserCommandBase):
    """运行日志核心指令"""

    def __init__(self, bot: Bot):
        super().__init__(bot)
        self.helper = _LogHelper(bot)
        self.messages = _LogMessages()
        self.log_usage = (
            "日志指令：\n"
            ".log new <名称>  创建并立即开始新的日志\n"
            ".log on [名称]  继续当前日志或开启指定日志\n"
            ".log off       暂停当前日志\n"
            ".log halt       停止当前日志（保留为未激活状态）\n"
            ".log end        停止并导出当前日志\n"
            ".log list       查看本群日志\n"
            ".log stat [名称] 查看日志统计\n"
            ".log get <名称>  查看最近一次导出信息\n"
            ".log del <名称>  删除指定日志\n"
            ".log set [选项]  切换过滤设置"
        )
        self.log_help = self.log_usage
        self.log_on_already = "日志已在进行中。"
        self.log_off_not_active = "当前没有正在记录的日志。"
        self.log_folder_policy = "说明：若群文件存在“跑团log”文件夹，导出的日志会上传至该文件夹；否则上传到群文件根目录。"
        self.log_folder_hint = "日志文件将上传到群文件夹（如存在）"\
            "跑团log"\
            "。"
        self.log_on_start = self.messages.new_started
        self.log_off_result = self.messages.end_summary
        self.bot.loc_helper.register_loc_text(LOC_LOG_SET_TOGGLED, "{item} 切换为 {state}", "切换日志设置")
        self.bot.cfg_helper.register_config(CFG_LOG_UPLOAD_ENABLE, "1", "是否在 .log end 时上传日志云端 (1/0)")
        self.bot.cfg_helper.register_config(CFG_LOG_UPLOAD_ENDPOINT, UPLOAD_ENDPOINT_DEFAULT, "日志云端上传接口地址")
        self.bot.cfg_helper.register_config(CFG_LOG_UPLOAD_TOKEN, "", "日志云端上传授权 Token，可留空")
        # 允许配置内存中保留的最大日志记录条数，超出自动丢弃最早的记录以避免 OOM
        self.bot.cfg_helper.register_config(CFG_LOG_MAX_RECORDS, str(LOG_MAX_RECORDS_DEFAULT), "单个日志在内存中保留的最大消息条数，超过将自动丢弃最早的记录；-1 为不限制（不建议长期开启）")

    def can_process_msg(self, msg_str: str, meta: MessageMetaData) -> Tuple[bool, bool, Any]:
        if not msg_str.startswith(".log"):
            return False, False, None
        args = msg_str.split()
        action = args[1] if len(args) > 1 else ""
        name = " ".join(args[2:]).strip()
        return True, False, (action.lower(), name)

    def process_msg(self, msg_str: str, meta: MessageMetaData, hint: Any) -> List[BotCommandBase]:
        if not meta.group_id:
            return []
        action, param = hint if isinstance(hint, tuple) else (hint, "")
        group_id = meta.group_id
        payload = self.helper.load_payload(group_id)
        filters = _ensure_filters(payload)
        logs = payload.setdefault(LOG_GROUP_LOGS, {})
        current_id = payload.get(LOG_GROUP_CURRENT, "")
        current_entry = logs.get(current_id)
        reply_port = GroupMessagePort(group_id)

        cmds: List[BotCommandBase]
        if action == "new":
            feedback = self._handle_new(payload, group_id, param)
            cmds = [BotSendMsgCommand(self.bot.account, feedback, [reply_port])]
        elif action == "on":
            feedback = self._handle_on(payload, group_id, param)
            cmds = [BotSendMsgCommand(self.bot.account, feedback, [reply_port])]
        elif action == "off":
            feedback = self._handle_off(payload, group_id)
            cmds = [BotSendMsgCommand(self.bot.account, feedback, [reply_port])]
        elif action in ("halt", "stop"):
            feedback = self._handle_halt(payload, group_id)
            cmds = [BotSendMsgCommand(self.bot.account, feedback, [reply_port])]
        elif action == "end":
            cmds = self._handle_end(payload, group_id)
        elif action == "list":
            feedback = self._handle_list(payload)
            cmds = [BotSendMsgCommand(self.bot.account, feedback, [reply_port])]
        elif action == "del":
            feedback = self._handle_delete(payload, group_id, param)
            cmds = [BotSendMsgCommand(self.bot.account, feedback, [reply_port])]
        elif action == "get":
            feedback = self._handle_get(payload, group_id, param)
            cmds = [BotSendMsgCommand(self.bot.account, feedback, [reply_port])]
        elif action == "stat":
            feedback = self._handle_stat(payload, group_id, param)
            cmds = [BotSendMsgCommand(self.bot.account, feedback, [reply_port])]
        elif action == "set":
            feedback = self._handle_set(payload, filters, param)
            cmds = [BotSendMsgCommand(self.bot.account, feedback, [reply_port])]
        elif action == "":
            cmds = [BotSendMsgCommand(self.bot.account, self.log_usage, [reply_port])]
        else:
            cmds = [BotSendMsgCommand(self.bot.account, self.log_usage, [reply_port])]

        self.helper.save_payload(group_id, payload)
        return cmds

    def _handle_new(self, payload: Dict[str, Any], group_id: str, name: str) -> str:
        name = name.strip()
        if not name:
            return "请提供日志名称，例如 .log new 旅团日志"
        current = payload.get(LOG_GROUP_CURRENT, "")
        if current:
            current_entry = payload[LOG_GROUP_LOGS].get(current)
            current_name = current_entry.get(LOG_KEY_NAME, current) if current_entry else current
            return self.messages.need_empty_current.format(name=current_name)
        name_index = payload.setdefault(LOG_GROUP_NAME_INDEX, {})
        if name.lower() in name_index:
            return "已存在同名日志，请更换名称。"
        now = _now_str()
        log_id = _generate_log_id()
        payload[LOG_GROUP_LOGS][log_id] = {
            LOG_KEY_NAME: name,
            LOG_KEY_CREATED_AT: now,
            LOG_KEY_UPDATED_AT: now,
            LOG_KEY_RECORDING: True,
            LOG_KEY_RECORDS: [],
            LOG_KEY_COLOR_MAP: {},
            LOG_KEY_STATS: _empty_stats(),
            LOG_KEY_SESSION_COUNT: 0,
            LOG_KEY_RECORD_BEGIN_AT: now,
            LOG_KEY_LAST_WARN: now,
            LOG_KEY_UPLOAD: {},
        }
        payload[LOG_GROUP_NAME_INDEX][name.lower()] = log_id
        payload[LOG_GROUP_CURRENT] = log_id
        # 同步到 DB（元数据）
        if get_connection and upsert_log:
            try:
                conn = get_connection()
                try:
                    filters = payload.get(LOG_GROUP_FILTERS, DEFAULT_FILTERS)
                    upsert_log(conn, {
                        "id": log_id,
                        "group_id": group_id,
                        "name": name,
                        "created_at": now,
                        "updated_at": now,
                        "recording": True,
                        "record_begin_at": now,
                        "last_warn": now,
                        "filter_outside": int(bool(filters.get(FILTER_OUTSIDE))),
                        "filter_command": int(bool(filters.get(FILTER_COMMAND))),
                        "filter_bot": int(bool(filters.get(FILTER_BOT))),
                        "filter_media": int(bool(filters.get(FILTER_MEDIA))),
                        "filter_forum_code": int(bool(filters.get(FILTER_FORUM_CODE))),
                        "upload_time": None,
                        "upload_file": None,
                        "upload_note": None,
                        "url": None,
                    })
                    conn.commit()
                finally:
                    conn.close()
            except Exception as e:
                dice_log(f"[LogDB] upsert new log error: {e}")
        return self.messages.new_started.format(name=name)

    def _handle_on(self, payload: Dict[str, Any], group_id: str, name: str) -> str:
        logs = payload.get(LOG_GROUP_LOGS, {})
        if name:
            target_id = _find_log_id_by_name(payload, name)
            if not target_id:
                return self.messages.no_target.format(name=name)
            if payload.get(LOG_GROUP_CURRENT) != target_id:
                current_id = payload.get(LOG_GROUP_CURRENT)
                if current_id and current_id in logs:
                    logs[current_id][LOG_KEY_RECORDING] = False
                payload[LOG_GROUP_CURRENT] = target_id
            entry = logs[target_id]
        else:
            current_id = payload.get(LOG_GROUP_CURRENT)
            if not current_id:
                return self.messages.no_current
            entry = logs.get(current_id)
            if not entry:
                return self.messages.no_current
            target_id = current_id

        if entry.get(LOG_KEY_RECORDING):
            return self.messages.already_running.format(name=entry.get(LOG_KEY_NAME, target_id))

        now = _now_str()
        entry[LOG_KEY_RECORDING] = True
        entry[LOG_KEY_SESSION_COUNT] = 0
        entry[LOG_KEY_RECORD_BEGIN_AT] = now
        entry[LOG_KEY_LAST_WARN] = now
        entry[LOG_KEY_UPDATED_AT] = now
        logs[target_id] = entry
        payload[LOG_GROUP_LOGS] = logs
        # DB 同步
        if get_connection and upsert_log:
            try:
                conn = get_connection()
                try:
                    filters = payload.get(LOG_GROUP_FILTERS, DEFAULT_FILTERS)
                    upsert_log(conn, {
                        "id": target_id,
                        "group_id": group_id,
                        "name": entry.get(LOG_KEY_NAME, target_id),
                        "created_at": entry.get(LOG_KEY_CREATED_AT, now),
                        "updated_at": now,
                        "recording": True,
                        "record_begin_at": entry.get(LOG_KEY_RECORD_BEGIN_AT, now),
                        "last_warn": entry.get(LOG_KEY_LAST_WARN, now),
                        "filter_outside": int(bool(filters.get(FILTER_OUTSIDE))),
                        "filter_command": int(bool(filters.get(FILTER_COMMAND))),
                        "filter_bot": int(bool(filters.get(FILTER_BOT))),
                        "filter_media": int(bool(filters.get(FILTER_MEDIA))),
                        "filter_forum_code": int(bool(filters.get(FILTER_FORUM_CODE))),
                        "upload_time": entry.get(LOG_KEY_UPLOAD, {}).get(LOG_KEY_UPLOAD_TIME),
                        "upload_file": entry.get(LOG_KEY_UPLOAD, {}).get(LOG_KEY_UPLOAD_FILE),
                        "upload_note": entry.get(LOG_KEY_UPLOAD, {}).get(LOG_KEY_UPLOAD_NOTE),
                        "url": entry.get(LOG_KEY_UPLOAD, {}).get("url"),
                    })
                    conn.commit()
                finally:
                    conn.close()
            except Exception as e:
                dice_log(f"[LogDB] upsert on error: {e}")
        return self.messages.resume.format(name=entry.get(LOG_KEY_NAME, target_id))

    def _handle_off(self, payload: Dict[str, Any], group_id: str) -> str:
        current_id = payload.get(LOG_GROUP_CURRENT, "")
        if not current_id:
            return self.messages.no_current
        entry = payload[LOG_GROUP_LOGS].get(current_id)
        if not entry:
            payload[LOG_GROUP_CURRENT] = ""
            return self.messages.no_current
        if not entry.get(LOG_KEY_RECORDING):
            return self.messages.paused.format(name=entry.get(LOG_KEY_NAME, current_id))
        entry[LOG_KEY_RECORDING] = False
        entry[LOG_KEY_UPDATED_AT] = _now_str()
        payload[LOG_GROUP_LOGS][current_id] = entry
        # DB 同步
        if get_connection and upsert_log:
            try:
                conn = get_connection()
                try:
                    filters = payload.get(LOG_GROUP_FILTERS, DEFAULT_FILTERS)
                    upsert_log(conn, {
                        "id": current_id,
                        "group_id": group_id,
                        "name": entry.get(LOG_KEY_NAME, current_id),
                        "created_at": entry.get(LOG_KEY_CREATED_AT, _now_str()),
                        "updated_at": entry.get(LOG_KEY_UPDATED_AT),
                        "recording": False,
                        "record_begin_at": entry.get(LOG_KEY_RECORD_BEGIN_AT),
                        "last_warn": entry.get(LOG_KEY_LAST_WARN),
                        "filter_outside": int(bool(filters.get(FILTER_OUTSIDE))),
                        "filter_command": int(bool(filters.get(FILTER_COMMAND))),
                        "filter_bot": int(bool(filters.get(FILTER_BOT))),
                        "filter_media": int(bool(filters.get(FILTER_MEDIA))),
                        "filter_forum_code": int(bool(filters.get(FILTER_FORUM_CODE))),
                        "upload_time": entry.get(LOG_KEY_UPLOAD, {}).get(LOG_KEY_UPLOAD_TIME),
                        "upload_file": entry.get(LOG_KEY_UPLOAD, {}).get(LOG_KEY_UPLOAD_FILE),
                        "upload_note": entry.get(LOG_KEY_UPLOAD, {}).get(LOG_KEY_UPLOAD_NOTE),
                        "url": entry.get(LOG_KEY_UPLOAD, {}).get("url"),
                    })
                    conn.commit()
                finally:
                    conn.close()
            except Exception as e:
                dice_log(f"[LogDB] upsert off error: {e}")
        return self.messages.paused.format(name=entry.get(LOG_KEY_NAME, current_id))

    def _handle_halt(self, payload: Dict[str, Any], group_id: str) -> str:
        current_id = payload.get(LOG_GROUP_CURRENT, "")
        if not current_id:
            return self.messages.no_current
        entry = payload[LOG_GROUP_LOGS].get(current_id)
        if not entry:
            payload[LOG_GROUP_CURRENT] = ""
            return self.messages.no_current
        entry[LOG_KEY_RECORDING] = False
        entry[LOG_KEY_UPDATED_AT] = _now_str()
        payload[LOG_GROUP_LOGS][current_id] = entry
        payload[LOG_GROUP_CURRENT] = ""
        # DB 同步
        if get_connection and upsert_log:
            try:
                conn = get_connection()
                try:
                    filters = payload.get(LOG_GROUP_FILTERS, DEFAULT_FILTERS)
                    upsert_log(conn, {
                        "id": current_id,
                        "group_id": group_id,
                        "name": entry.get(LOG_KEY_NAME, current_id),
                        "created_at": entry.get(LOG_KEY_CREATED_AT, _now_str()),
                        "updated_at": entry.get(LOG_KEY_UPDATED_AT),
                        "recording": False,
                        "record_begin_at": entry.get(LOG_KEY_RECORD_BEGIN_AT),
                        "last_warn": entry.get(LOG_KEY_LAST_WARN),
                        "filter_outside": int(bool(filters.get(FILTER_OUTSIDE))),
                        "filter_command": int(bool(filters.get(FILTER_COMMAND))),
                        "filter_bot": int(bool(filters.get(FILTER_BOT))),
                        "filter_media": int(bool(filters.get(FILTER_MEDIA))),
                        "filter_forum_code": int(bool(filters.get(FILTER_FORUM_CODE))),
                        "upload_time": entry.get(LOG_KEY_UPLOAD, {}).get(LOG_KEY_UPLOAD_TIME),
                        "upload_file": entry.get(LOG_KEY_UPLOAD, {}).get(LOG_KEY_UPLOAD_FILE),
                        "upload_note": entry.get(LOG_KEY_UPLOAD, {}).get(LOG_KEY_UPLOAD_NOTE),
                        "url": entry.get(LOG_KEY_UPLOAD, {}).get("url"),
                    })
                    conn.commit()
                finally:
                    conn.close()
            except Exception as e:
                dice_log(f"[LogDB] upsert halt error: {e}")
        return self.messages.halted.format(name=entry.get(LOG_KEY_NAME, current_id))

    def _handle_end(self, payload: Dict[str, Any], group_id: str) -> List[BotCommandBase]:
        current_id = payload.get(LOG_GROUP_CURRENT, "")
        if not current_id:
            feedback = self.messages.no_current
            return [BotSendMsgCommand(self.bot.account, feedback, [GroupMessagePort(group_id)])]
        entry = payload[LOG_GROUP_LOGS].get(current_id)
        if not entry:
            payload[LOG_GROUP_CURRENT] = ""
            feedback = self.messages.no_current
            return [BotSendMsgCommand(self.bot.account, feedback, [GroupMessagePort(group_id)])]
        entry[LOG_KEY_RECORDING] = False
        entry[LOG_KEY_UPDATED_AT] = _now_str()
        payload[LOG_GROUP_CURRENT] = ""
        payload[LOG_GROUP_LOGS][current_id] = entry

        filters = _ensure_filters(payload)
        file_main_path, display_name, extra_files = self._generate_file(group_id, entry, filters, log_id=current_id)
        count = entry.get(LOG_KEY_STATS, {}).get("messages", len(entry.get(LOG_KEY_RECORDS, [])))
        upload_feedback = self._try_upload_log(group_id, entry, log_id=current_id)
        upload_note = "上传至群文件"
        upload_url = None
        if upload_feedback:
            upload_note = upload_feedback.get("message", upload_note)
            upload_url = upload_feedback.get("url")
        entry[LOG_KEY_UPLOAD] = {
            LOG_KEY_UPLOAD_TIME: _now_str(),
            LOG_KEY_UPLOAD_FILE: display_name,
            LOG_KEY_UPLOAD_NOTE: upload_note,
        }
        if upload_url:
            entry[LOG_KEY_UPLOAD]["url"] = upload_url
        payload[LOG_GROUP_LOGS][current_id] = entry
        # DB 更新上传信息
        if get_connection and update_log_upload:
            try:
                conn = get_connection()
                try:
                    update_log_upload(conn, current_id, {
                        "time": entry[LOG_KEY_UPLOAD].get(LOG_KEY_UPLOAD_TIME),
                        "file": entry[LOG_KEY_UPLOAD].get(LOG_KEY_UPLOAD_FILE),
                        "note": entry[LOG_KEY_UPLOAD].get(LOG_KEY_UPLOAD_NOTE),
                        "url": entry[LOG_KEY_UPLOAD].get("url"),
                    })
                    conn.commit()
                finally:
                    conn.close()
            except Exception as e:
                dice_log(f"[LogDB] update upload error: {e}")

        feedback_lines = [self.messages.end_summary.format(name=entry.get(LOG_KEY_NAME, current_id), count=count)]
        if upload_feedback:
            if upload_feedback.get("success") and upload_url:
                feedback_lines.append(f"线上日志链接：{upload_url}")
            elif upload_feedback.get("message"):
                feedback_lines.append(f"云端上传提示：{upload_feedback['message']}")
        feedback = "\n".join(feedback_lines)
        port = GroupMessagePort(group_id)
        folder_prefix = "跑团log/"
        commands: List[BotCommandBase] = [
            BotSendMsgCommand(self.bot.account, feedback, [port]),
            BotSendFileCommand(self.bot.account, file_main_path, folder_prefix + display_name, [port]),
        ]
        for fpath, fname in extra_files:
            commands.append(BotSendFileCommand(self.bot.account, fpath, folder_prefix + fname, [port]))
        return commands

    def _handle_list(self, payload: Dict[str, Any]) -> str:
        logs = payload.get(LOG_GROUP_LOGS, {})
        if not logs:
            return "当前没有保存任何日志。"
        current_id = payload.get(LOG_GROUP_CURRENT, "")
        lines = [_LogFormatter.status_line(log_id, entry, current_id) for log_id, entry in logs.items()]
        return "\n".join([self.messages.list_header] + lines)

    def _handle_delete(self, payload: Dict[str, Any], group_id: str, name: str) -> str:
        if not name:
            return "请提供要删除的日志名称。"
        log_id = _find_log_id_by_name(payload, name)
        if not log_id:
            return self.messages.no_target.format(name=name)
        current_id = payload.get(LOG_GROUP_CURRENT)
        entry = payload[LOG_GROUP_LOGS].get(log_id)
        if entry and entry.get(LOG_KEY_RECORDING):
            return self.messages.pause_before_delete.format(name=entry.get(LOG_KEY_NAME, name))
        if current_id == log_id:
            payload[LOG_GROUP_CURRENT] = ""
        payload[LOG_GROUP_LOGS].pop(log_id, None)
        payload[LOG_GROUP_NAME_INDEX] = {
            k: v for k, v in payload.get(LOG_GROUP_NAME_INDEX, {}).items() if v != log_id
        }
        # DB 删除（级联删除记录）
        if get_connection and delete_log:
            try:
                conn = get_connection()
                try:
                    delete_log(conn, log_id)
                    conn.commit()
                finally:
                    conn.close()
            except Exception as e:
                dice_log(f"[LogDB] delete log error: {e}")
        return self.messages.deleted.format(name=entry.get(LOG_KEY_NAME, name) if entry else name)

    def _handle_get(self, payload: Dict[str, Any], group_id: str, name: str) -> str:
        if not name:
            return "请提供日志名称。"
        log_id = _find_log_id_by_name(payload, name)
        if not log_id:
            return self.messages.no_target.format(name=name)
        entry = payload[LOG_GROUP_LOGS].get(log_id)
        if not entry:
            return self.messages.no_target.format(name=name)
        upload = entry.get(LOG_KEY_UPLOAD, {})
        should_retry = not upload or not upload.get("url")
        retry_feedback: Optional[Dict[str, Any]] = None
        if should_retry:
            retry_feedback = self._try_upload_log(group_id, entry, log_id=log_id)
            if retry_feedback:
                upload_file_name = upload.get(LOG_KEY_UPLOAD_FILE) or f"{entry.get(LOG_KEY_NAME, name)}"
                note = retry_feedback.get("message", "")
                entry[LOG_KEY_UPLOAD] = {
                    LOG_KEY_UPLOAD_TIME: _now_str(),
                    LOG_KEY_UPLOAD_FILE: upload_file_name,
                    LOG_KEY_UPLOAD_NOTE: note or "自动上传",
                }
                if retry_feedback.get("url"):
                    entry[LOG_KEY_UPLOAD]["url"] = retry_feedback["url"]
                payload[LOG_GROUP_LOGS][log_id] = entry
                upload = entry.get(LOG_KEY_UPLOAD, {})
        if not upload:
            message = "该日志尚未导出过。"
            if retry_feedback and retry_feedback.get("message"):
                message += f"\n云端上传提示：{retry_feedback['message']}"
            return message
        upload_time = upload.get(LOG_KEY_UPLOAD_TIME, "未知时间")
        upload_file = upload.get(LOG_KEY_UPLOAD_FILE, "未知文件")
        url = upload.get("url")
        lines = [f"日志《{entry.get(LOG_KEY_NAME, name)}》最近一次导出：{upload_time}，文件名：{upload_file}"]
        if url:
            lines.append(f"线上查看：{url}")
        elif retry_feedback and retry_feedback.get("message"):
            lines.append(f"云端上传提示：{retry_feedback['message']}")
        return "\n".join(lines)

    def _handle_stat(self, payload: Dict[str, Any], group_id: str, name: str) -> str:
        logs = payload.get(LOG_GROUP_LOGS, {})
        target_entry: Optional[Dict[str, Any]] = None
        if name:
            log_id = _find_log_id_by_name(payload, name)
            if log_id:
                target_entry = logs.get(log_id)
        else:
            current_id = payload.get(LOG_GROUP_CURRENT)
            if current_id:
                target_entry = logs.get(current_id)
        if not target_entry:
            return "未找到对应日志，或当前没有正在使用的日志。"
        return _StatsFormatter.format(target_entry)

    def _handle_set(self, payload: Dict[str, Any], filters: Dict[str, bool], param: str) -> str:
        if not param:
            status_lines = [
                f"场外发言过滤：{'ON' if filters.get(FILTER_OUTSIDE) else 'OFF'}",
                f"指令过滤：{'ON' if filters.get(FILTER_COMMAND) else 'OFF'}",
                f"bot 过滤：{'ON' if filters.get(FILTER_BOT) else 'OFF'}",
                f"图片表情过滤：{'ON' if filters.get(FILTER_MEDIA) else 'OFF'}",
                f"论坛代码生成：{'ON' if filters.get(FILTER_FORUM_CODE) else 'OFF'}",
            ]
            return "日志过滤设置：\n" + "\n".join(status_lines)
        alias_map = {
            "outside": FILTER_OUTSIDE,
            "场外": FILTER_OUTSIDE,
            "场外发言过滤": FILTER_OUTSIDE,
            "command": FILTER_COMMAND,
            "指令": FILTER_COMMAND,
            "指令过滤": FILTER_COMMAND,
            "bot": FILTER_BOT,
            "bot过滤": FILTER_BOT,
            "media": FILTER_MEDIA,
            "图片": FILTER_MEDIA,
            "图片表情过滤": FILTER_MEDIA,
            "forum_code": FILTER_FORUM_CODE,
            "论坛": FILTER_FORUM_CODE,
            "论坛代码": FILTER_FORUM_CODE,
            "论坛代码生成": FILTER_FORUM_CODE,
        }
        key = alias_map.get(param.lower())
        if not key:
            return "未知选项，可用：outside / command / bot / media / forum_code"
        current_value = filters.get(key, False)
        filters[key] = not current_value
        state = "ON" if filters[key] else "OFF"
        return self.bot.loc_helper.format_loc_text(LOC_LOG_SET_TOGGLED, item=param, state=state)

    def _generate_file(self, group_id: str, log_entry: Dict[str, Any], filters: Dict[str, bool], *, log_id: Optional[str] = None) -> Tuple[str, str, List[Tuple[str, str]]]:
        # 优先从 DB 读取记录，避免占用内存
        records: List[Dict[str, Any]]
        # Prefer DB records but merge with any legacy in-memory records so that
        # old-format logs (stored in payload) and new DB-backed records are
        # concatenated for export.
        records_payload = list(log_entry.get(LOG_KEY_RECORDS, []))
        records_db = []
        if get_connection and fetch_records:
            try:
                conn = get_connection()
                try:
                    use_log_id = log_id or self._get_log_id_by_entry(group_id, log_entry)
                    records_db = fetch_records(conn, use_log_id)
                finally:
                    conn.close()
            except Exception as e:
                dice_log(f"[LogDB] fetch_records error: {e}")
                records_db = []
        # Merge DB-backed and payload records, sort by timestamp to preserve order
        if records_db and records_payload:
            try:
                combined = records_db + records_payload
                combined.sort(key=lambda r: str_to_datetime(r.get('time', _now_str())) or _now_str())
                records = combined
            except Exception:
                records = records_db + records_payload
        elif records_db:
            records = records_db
        else:
            records = records_payload
        color_map = dict(log_entry.get(LOG_KEY_COLOR_MAP, {}))
        log_name = log_entry.get(LOG_KEY_NAME, "log")
        start_time = log_entry.get(LOG_KEY_CREATED_AT, _now_str())
        safe_name = _sanitize_filename(log_name)
        safe_start = start_time.replace('/', '-').replace(':', '-').replace(' ', '_')
        display_name_base = f"{safe_name}_{safe_start}"
        logs_dir = os.path.join(self.bot.data_path, "logs")
        os.makedirs(logs_dir, exist_ok=True)

        nickname_cache: Dict[str, str] = {}
        for record in records:
            uid = record.get('user_id')
            if uid and uid not in nickname_cache:
                try:
                    nick = self.bot.get_nickname(uid, group_id)
                except Exception:
                    nick = None
                if nick and nick not in ("UNDEF_NAME", "----"):
                    nickname_cache[uid] = nick

        msg_map: Dict[str, Dict[str, str]] = {}
        for record in records:
            mid = record.get('message_id')
            if not mid:
                continue
            raw_content = record.get('content', '')
            cleaned = re.sub(r"\[CQ:reply,(?:id|reply|source_id)=\d+[^\]]*\]", "", raw_content)
            sender_display = nickname_cache.get(record.get('user_id')) or record.get('nickname', '?')
            msg_map[str(mid)] = {"content": cleaned, "nickname": sender_display}

        user_display: Dict[str, str] = {}
        for record in records:
            uid = record.get('user_id')
            if uid and uid not in user_display:
                user_display[uid] = nickname_cache.get(uid) or record.get('nickname') or uid

        def humanize_cq(raw: str) -> str:
            text = raw

            def repl_reply(match: re.Match) -> str:
                rid = match.group(1)
                origin = msg_map.get(rid)
                if not origin:
                    return "| 引用消息不在 log 范围内\n"
                origin_content = origin['content'].strip() or "(空白)"
                lines = [ln.strip() for ln in origin_content.splitlines() if ln.strip()][:3] or [origin_content]
                lines = [ln[:60] + ('…' if len(ln) > 60 else '') for ln in lines]
                quote_lines = [f"| {origin['nickname']}"] + [f"| {ln}" for ln in lines]
                return "\n".join(quote_lines) + "\n"

            text = re.sub(r"\[CQ:reply,(?:id|reply|source_id)=(\d+)[^\]]*\]", repl_reply, text)

            def repl_at(match: re.Match) -> str:
                uid = match.group(1)
                nick = user_display.get(uid)
                if not nick or nick in ("UNDEF_NAME", "----"):
                    try:
                        nick = self.bot.get_nickname(uid, group_id) or uid
                    except Exception:
                        nick = uid
                return f"@{nick}"

            text = re.sub(r"\[CQ:at,qq=(\d+)(?:,[^\]]*)?\]", repl_at, text)
            return text

        txt_path = os.path.join(logs_dir, display_name_base + ".txt")
        with open(txt_path, "w", encoding="utf-8") as txt_file:
            txt_file.write(f"群 {group_id} 跑团日志 (开始于 {start_time})\n\n")
            for record in records:
                uid = record.get('user_id', '?')
                display_name = nickname_cache.get(uid) or record.get('nickname') or ("骰娘" if uid == self.bot.account else uid)
                color_map.setdefault(uid, _pick_color(color_map, uid))
                content_out = humanize_cq(record.get('content', ''))
                txt_file.write(f"{display_name} ({uid})  {record.get('time', '?')}\n")
                txt_file.write(content_out + "\n\n")

        docx_path = None
        try:
            from docx import Document  # type: ignore
            from docx.shared import RGBColor  # type: ignore

            document = Document()
            document.add_heading(f"群 {group_id} 跑团日志 (开始于 {start_time})", level=1)
            for record in records:
                uid = record.get('user_id', '?')
                color_hex = color_map.get(uid, "000000")
                r = int(color_hex[0:2], 16)
                g = int(color_hex[2:4], 16)
                b = int(color_hex[4:6], 16)
                display_name = nickname_cache.get(uid) or record.get('nickname') or ("骰娘" if uid == self.bot.account else uid)
                body = document.add_paragraph()
                run_body = body.add_run(f"<{display_name}>{humanize_cq(record.get('content', ''))}")
                run_body.font.color.rgb = RGBColor(r, g, b)
            docx_path = os.path.join(logs_dir, display_name_base + ".docx")
            document.save(docx_path)
        except Exception as exc:
            dice_log(f"[LogExport] docx generation failed: {type(exc).__name__}: {exc}")
            docx_path = None

        extra_files: List[Tuple[str, str]] = []
        if filters.get(FILTER_FORUM_CODE):
            forum_txt_path = os.path.join(logs_dir, display_name_base + "_forum.txt")
            try:
                with open(forum_txt_path, "w", encoding="utf-8") as forum_file:
                    forum_file.write(_generate_forum_code_log(records))
                extra_files.append((forum_txt_path, os.path.basename(forum_txt_path)))
            except Exception as exc:
                dice_log(f"[LogExport] forum code generation failed: {type(exc).__name__}: {exc}")
        if docx_path:
            extra_files.append((txt_path, os.path.basename(txt_path)))
            return docx_path, os.path.basename(docx_path), extra_files
        return txt_path, os.path.basename(txt_path), extra_files

    def _get_upload_settings(self) -> Dict[str, Any]:
        try:
            enabled = self.bot.cfg_helper.get_config(CFG_LOG_UPLOAD_ENABLE)[0]
        except Exception:
            enabled = "1"
        try:
            endpoint = self.bot.cfg_helper.get_config(CFG_LOG_UPLOAD_ENDPOINT)[0]
        except Exception:
            endpoint = UPLOAD_ENDPOINT_DEFAULT
        try:
            token = self.bot.cfg_helper.get_config(CFG_LOG_UPLOAD_TOKEN)[0]
        except Exception:
            token = ""
        endpoint = endpoint.strip() or UPLOAD_ENDPOINT_DEFAULT
        return {
            "enabled": str(enabled).strip() != "0",
            "endpoint": endpoint,
            "token": token.strip(),
        }

    def _build_upload_payload(self, log_entry: Dict[str, Any], log_id: Optional[str] = None) -> Optional[Dict[str, Any]]:
        # 尝试从 DB 获取记录
        # Merge DB and payload records for upload payload as well
        records_payload = list(log_entry.get(LOG_KEY_RECORDS, []))
        records_db = []
        if get_connection and fetch_records:
            try:
                conn = get_connection()
                try:
                    use_log_id = log_id or self._get_log_id_by_entry(group_id, log_entry)
                    records_db = fetch_records(conn, use_log_id)
                finally:
                    conn.close()
            except Exception as e:
                dice_log(f"[LogDB] fetch_records for upload error: {e}")
                records_db = []
        if records_db and records_payload:
            try:
                combined = records_db + records_payload
                combined.sort(key=lambda r: str_to_datetime(r.get('time', _now_str())) or _now_str())
                records = combined
            except Exception:
                records = records_db + records_payload
        elif records_db:
            records = records_db
        else:
            records = records_payload
        if not records:
            return None
        items: List[Dict[str, Any]] = []
        for record in records:
            raw_uid = record.get("user_id") or record.get("imUserId") or ""
            if isinstance(raw_uid, int):
                raw_uid = str(raw_uid)
            if not raw_uid and record.get("uniformId"):
                raw_uid = str(record.get("uniformId")).split(":")[-1]
            user_id = str(raw_uid or "")
            nickname = record.get("nickname", user_id or "?")
            try:
                timestamp = int(str_to_datetime(record.get("time", _now_str())).timestamp())
            except Exception:
                timestamp = int(time.time())
            content = record.get("content", "")
            is_bot_msg = record.get(LOG_KEY_SOURCE) == "bot"
            roll_result = _detect_roll_result(content) if is_bot_msg else None
            is_dice = bool(roll_result)
            command_info: Optional[Dict[str, Any]] = None
            if roll_result:
                command_info = {
                    "cmd": "roll",
                    "result": roll_result,
                }
            items.append({
                "nickname": nickname,
                "imUserId": user_id,
                "uniformId": f"QQ:{user_id}" if user_id else "",
                "time": timestamp,
                "message": content,
                "isDice": is_dice,
                "commandId": record.get("message_id") or "",
                "commandInfo": command_info,
                "rawMsgId": record.get("message_id") or "",
            })
        payload = {
            "version": UPLOAD_VERSION,
            "items": items,
        }
        json_str = json.dumps(payload, ensure_ascii=False)
        compressed = zlib.compress(json_str.encode("utf-8"))
        return {
            "file": compressed,
            "name": log_entry.get(LOG_KEY_NAME, "日志"),
        }

    def _try_upload_log(self, group_id: str, log_entry: Dict[str, Any], *, log_id: Optional[str] = None) -> Optional[Dict[str, Any]]:
        settings = self._get_upload_settings()
        if not settings.get("enabled"):
            return None
        if requests is None:
            return {"success": False, "message": "requests 模块不可用，已跳过云端上传"}
        # 获取当前日志 ID 以便从 DB 读取
        use_log_id = log_id or self._get_log_id_by_entry(group_id, log_entry)
        payload_data = self._build_upload_payload(log_entry, log_id=use_log_id)
        if not payload_data:
            return {"success": False, "message": "日志内容为空，已跳过云端上传"}
        files = {
            'file': ('log-zlib-compressed', payload_data['file'], 'application/octet-stream')
        }
        try:
            masters = self.bot.cfg_helper.get_config(CFG_MASTER)
        except Exception:
            masters = []
        uploader_id = None
        for mid in masters:
            if mid.strip():
                uploader_id = mid.strip()
                break
        if not uploader_id:
            uploader_id = str(self.bot.account)
        form = {
            'name': payload_data['name'],
            'uniform_id': f"QQ:{uploader_id}",
            'client': 'DicePP',
            'version': str(UPLOAD_VERSION),
        }
        headers = {}
        if settings.get("token"):
            headers['Authorization'] = f"Bearer {settings['token']}"
        try:
            response = requests.put(settings['endpoint'], data=form, files=files, headers=headers, timeout=15)
        except Exception as exc:
            return {"success": False, "message": f"云端上传失败：{exc}"}
        try:
            resp_json = response.json()
        except Exception:
            resp_json = {}
        if response.ok and isinstance(resp_json, dict) and resp_json.get('url'):
            return {
                "success": True,
                "message": "云端上传成功",
                "url": resp_json['url'],
            }
        msg = resp_json.get('message') if isinstance(resp_json, dict) else response.text
        return {
            "success": False,
            "message": f"云端上传失败：HTTP {response.status_code} {msg or response.reason}",
        }

    def get_help(self, keyword: str, meta: MessageMetaData) -> str:
        if keyword in ("log", "日志"):
            return self.log_usage
        return ""

    def get_description(self) -> str:
        return ".log 日志管理"

    # 辅助：通过 log_entry 反查其在 payload 中的 log_id（调用点都在命令处理期，可安全扫描一次）
    def _get_log_id_by_entry(self, group_id: str, entry: Dict[str, Any]) -> str:
        try:
            payload = _load_group_payload(self.bot, group_id) if group_id else None
            logs = payload.get(LOG_GROUP_LOGS, {}) if isinstance(payload, dict) else {}
            for lid, ent in logs.items():
                if ent is entry:
                    return lid
        except Exception:
            pass
        # 兜底：无法定位则退回名称
        return entry.get(LOG_KEY_NAME, "unknown")


# 提供给适配器：按消息撤回删除对应 DB 记录
def delete_log_record_by_message_id(bot: Bot, group_id: str, message_id: str) -> None:
    try:
        payload = _load_group_payload(bot, group_id)
        current_id = payload.get(LOG_GROUP_CURRENT, "")
        if not current_id:
            return
        if get_connection and delete_records_by_message_id:
            conn = get_connection()
            try:
                delete_records_by_message_id(conn, current_id, str(message_id))
                conn.commit()
            finally:
                conn.close()
    except Exception as e:
        try:
            dice_log(f"[LogDB] delete by message_id error: {e}")
        except Exception:
            pass


@custom_user_command(readable_name="跑团日志记录器", priority=DPP_COMMAND_PRIORITY_USUAL_LOWER_BOUND - 10,
                     flag=0, cluster=DPP_COMMAND_CLUSTER_DEFAULT, group_only=True)
class LogRecorderCommand(UserCommandBase):
    def __init__(self, bot: Bot):
        super().__init__(bot)
        self.helper = _LogHelper(bot)

    def can_process_msg(self, msg_str: str, meta: MessageMetaData) -> Tuple[bool, bool, Any]:
        if not meta.group_id:
            return False, False, None
        payload = _load_group_payload(self.bot, meta.group_id)
        current_id = payload.get(LOG_GROUP_CURRENT, "")
        if not current_id:
            return False, False, None
        entry = payload.get(LOG_GROUP_LOGS, {}).get(current_id)
        if not entry or not entry.get(LOG_KEY_RECORDING):
            return False, False, None
        return True, True, current_id

    def process_msg(self, msg_str: str, meta: MessageMetaData, hint: Any) -> List[BotCommandBase]:
        group_id = meta.group_id
        if not group_id:
            return []
        raw_content = getattr(meta, 'raw_msg', '') or msg_str
        return record_incoming_message(
            self.bot,
            group_id,
            str(meta.user_id or ""),
            meta.nickname or meta.user_id,
            raw_content,
            getattr(meta, 'message_id', None),
            is_bot=False,
        )

    def get_help(self, keyword: str, meta: MessageMetaData) -> str:
        return ""

    def get_description(self) -> str:
        return ""


@custom_user_command(readable_name="日志统计指令", priority=DPP_COMMAND_PRIORITY_DEFAULT,
                     flag=DPP_COMMAND_FLAG_INFO, cluster=DPP_COMMAND_CLUSTER_DEFAULT, group_only=True)
class LogStatCommand(UserCommandBase):
    def __init__(self, bot: Bot):
        super().__init__(bot)
        self.helper = _LogHelper(bot)

    def can_process_msg(self, msg_str: str, meta: MessageMetaData) -> Tuple[bool, bool, Any]:
        if not msg_str.startswith(".stat"):
            return False, False, None
        args = msg_str.split()
        if len(args) >= 2 and args[1] in ("log", "日志"):
            name = " ".join(args[2:]).strip()
            return True, False, name
        return False, False, None

    def process_msg(self, msg_str: str, meta: MessageMetaData, hint: Any) -> List[BotCommandBase]:
        if not meta.group_id:
            return []
        payload = _load_group_payload(self.bot, meta.group_id)
        logs = payload.get(LOG_GROUP_LOGS, {})
        name = hint if isinstance(hint, str) else ""
        entry: Optional[Dict[str, Any]] = None
        if name:
            log_id = _find_log_id_by_name(payload, name)
            if log_id:
                entry = logs.get(log_id)
        else:
            current_id = payload.get(LOG_GROUP_CURRENT, "")
            if current_id:
                entry = logs.get(current_id)
        if not entry:
            feedback = "未找到对应日志，或当前没有正在使用的日志。"
        else:
            feedback = _StatsFormatter.format(entry)
        return [BotSendMsgCommand(self.bot.account, feedback, [GroupMessagePort(meta.group_id)])]

    def get_help(self, keyword: str, meta: MessageMetaData) -> str:
        if keyword in ("stat", "统计"):
            return ".stat log [日志名] 查看指定日志的统计信息"
        return ""

    def get_description(self) -> str:
        return ".stat log 查看日志统计"
