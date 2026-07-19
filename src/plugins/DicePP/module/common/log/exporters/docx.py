from __future__ import annotations

from hashlib import sha256
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from ..projection import LogProjection
from ..types import LogExportFormat
from .base import AtomicFileExporter
from .metadata import export_view_label


USER_COLOR_PALETTE = (
    "C62828",
    "1565C0",
    "2E7D32",
    "EF6C00",
    "6A1B9A",
    "AD1457",
    "00838F",
    "4E342E",
    "283593",
    "00695C",
    "9E9D24",
    "4527A0",
)


class DocxLogExporter(AtomicFileExporter):
    format = LogExportFormat.DOCX
    suffix = ".docx"

    def _write_sync(self, projection: LogProjection, path: Path) -> None:
        document = Document()
        document.add_heading(projection.log_name, level=1)
        metadata = document.add_paragraph()
        metadata.add_run(
            f"日志 ID：{projection.log_id}\n"
            f"群 ID：{projection.group_id}\n"
            f"创建时间：{projection.created_at:%Y-%m-%d %H:%M:%S}\n"
            f"导出视图：{export_view_label(projection.view)}\n"
            f"记录快照：{projection.record_upper_id}"
        )

        for message in projection.messages:
            header = document.add_paragraph()
            name_run = header.add_run(f"{message.nickname} ({message.user_id})")
            name_run.bold = True
            name_run.font.color.rgb = stable_user_color(message.user_id)
            time_run = header.add_run(f"  {message.time:%Y-%m-%d %H:%M:%S}")
            time_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

            if message.reply is not None:
                reply = document.add_paragraph()
                reply.paragraph_format.left_indent = Pt(18)
                reply_run = reply.add_run(
                    f"回复消息：{message.reply.message_id}"
                    if message.reply.author is None
                    else f"{message.reply.author}（消息 {message.reply.message_id}）"
                )
                reply_run.italic = True
                reply_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                for line in message.reply.excerpt:
                    reply_run.add_break(WD_BREAK.LINE)
                    reply_run.add_text(line)

            body = document.add_paragraph(message.readable_text)
            body.paragraph_format.space_after = Pt(8)

        _apply_fonts(document)
        document.save(path)


def stable_user_color(user_id: str) -> RGBColor:
    digest = sha256(user_id.encode("utf-8")).digest()
    value = USER_COLOR_PALETTE[int.from_bytes(digest[:4], "big") % len(USER_COLOR_PALETTE)]
    return RGBColor.from_string(value)


def _apply_fonts(document: Document) -> None:
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Microsoft YaHei"
            run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "微软雅黑")
