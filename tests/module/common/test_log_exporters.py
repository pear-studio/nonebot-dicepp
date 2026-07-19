from __future__ import annotations

from datetime import datetime
from pathlib import Path

import pytest
from docx import Document

from module.common.log import (
    LogExportFormat,
    LogExportView,
    LogProjection,
    ProjectedMessage,
    ProjectedPart,
    ProjectedReply,
)
from module.common.log.exporters import (
    DocxLogExporter,
    TextLogExporter,
    build_filename_base,
    reserve_export_target,
    sanitize_filename_component,
)
from module.common.log.exporters.docx import stable_user_color

pytestmark = [pytest.mark.integration, pytest.mark.log]

NOW = datetime(2026, 7, 20, 16, 0, 0)


def _projection(view: LogExportView = LogExportView.CURATED) -> LogProjection:
    return LogProjection(
        log_id="log-12345678",
        group_id="group-1",
        log_name="雾都夜话",
        created_at=NOW,
        view=view,
        record_upper_id=2,
        messages=(
            ProjectedMessage(
                record_id=1,
                time=NOW,
                user_id="user-1",
                nickname="调查员",
                source="user",
                message_type="ambient",
                reply=None,
                parts=(ProjectedPart("text", "打开房门"),),
            ),
            ProjectedMessage(
                record_id=2,
                time=NOW,
                user_id="user-2",
                nickname="助手",
                source="bot",
                message_type="command",
                reply=ProjectedReply("m1", "调查员", ("打开房门",)),
                parts=(
                    ProjectedPart("text", "门后有地图"),
                    ProjectedPart("image", "[图片未归档]"),
                ),
            ),
        ),
    )


def _target(tmp_path: Path, format: LogExportFormat, suffix: str):
    data_root = tmp_path / "bot"
    return reserve_export_target(
        output_root=data_root / "logs",
        bot_data_root=data_root,
        filename_base="雾都夜话_群group-1_log-1234_req-1234_20260720-160000",
        request_id="req-12345678",
        format=format,
        suffix=suffix,
    )


@pytest.mark.asyncio
@pytest.mark.parametrize(
    ("view", "view_label"),
    [
        (LogExportView.CURATED, "跑团正文"),
        (LogExportView.COMPLETE, "全部记录"),
    ],
)
async def test_text_and_docx_export_the_same_projection_semantics(
    tmp_path: Path,
    view: LogExportView,
    view_label: str,
) -> None:
    projection = _projection(view)
    text_target = _target(tmp_path, LogExportFormat.TXT, ".txt")
    docx_target = _target(tmp_path, LogExportFormat.DOCX, ".docx")

    text_artifact = await TextLogExporter().generate(projection, text_target)
    docx_artifact = await DocxLogExporter().generate(projection, docx_target)

    text = text_artifact.path.read_text(encoding="utf-8")
    assert "日志：雾都夜话" in text
    assert "日志 ID：log-12345678" in text
    assert "群 ID：group-1" in text
    assert f"导出视图：{view_label}" in text
    assert "> 调查员（消息 m1）" in text
    assert "门后有地图[图片未归档]" in text
    document = Document(docx_artifact.path)
    document_text = "\n".join(p.text for p in document.paragraphs)
    for expected in (
        "雾都夜话",
        "日志 ID：log-12345678",
        f"导出视图：{view_label}",
        "打开房门",
        "门后有地图[图片未归档]",
    ):
        assert expected in document_text
    assert text_artifact.db_local_path.startswith("logs/")
    assert docx_artifact.size > 0
    assert not list((tmp_path / "bot" / "logs").glob("*.tmp"))

    reply_paragraph = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith("调查员（消息 m1）")
    )
    assert reply_paragraph.paragraph_format.left_indent is not None
    assert reply_paragraph.runs[0].italic is True
    name_paragraph = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith("调查员 (user-1)")
    )
    assert name_paragraph.runs[0].font.color.rgb == stable_user_color("user-1")


def test_user_color_is_stable_and_filename_rules_do_not_overwrite(tmp_path: Path) -> None:
    assert stable_user_color("same-user") == stable_user_color("same-user")
    assert sanitize_filename_component("CON") == "_CON"
    assert sanitize_filename_component(' bad:name?. ') == "bad_name_"
    assert len(sanitize_filename_component("x" * 200, maximum=20)) == 20

    data_root = tmp_path / "bot"
    base = build_filename_base(
        log_name='夜/话:*?"<>|.',
        group_id="group-1",
        log_id="log-123456789",
        request_id="request-123456789",
        timestamp="20260720-160000",
    )
    first = reserve_export_target(
        output_root=data_root / "logs",
        bot_data_root=data_root,
        filename_base=base,
        request_id="request-1",
        format=LogExportFormat.TXT,
        suffix=".txt",
    )
    first.final_path.write_text("keep", encoding="utf-8")
    second = reserve_export_target(
        output_root=data_root / "logs",
        bot_data_root=data_root,
        filename_base=base,
        request_id="request-2",
        format=LogExportFormat.TXT,
        suffix=".txt",
    )

    assert first.final_path != second.final_path
    assert first.final_path.read_text(encoding="utf-8") == "keep"
    assert second.group_file_name.endswith("_2.txt")
    assert not any(char in base for char in '<>:"/\\|?*')


def test_escaping_output_root_has_no_filesystem_side_effect(tmp_path: Path) -> None:
    data_root = tmp_path / "bot"
    escaping_root = tmp_path / "outside" / "logs"

    with pytest.raises(ValueError, match="escapes"):
        reserve_export_target(
            output_root=escaping_root,
            bot_data_root=data_root,
            filename_base="log",
            request_id="request",
            format=LogExportFormat.TXT,
            suffix=".txt",
        )

    assert not escaping_root.exists()
